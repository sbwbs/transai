#!/usr/bin/env python3
"""
Clinical Trial Protocol Term Extractor
Extracts and aligns Korean-English terminology from protocol documents
"""

import json
import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from collections import Counter
from docx import Document
import pandas as pd


class ProtocolTermExtractor:
    """Extract and analyze clinical trial protocol terminology"""
    
    def __init__(self):
        self.core_terms = {
            # Core protocol terminology
            '임상시험': 'clinical trial',
            '임상연구': 'clinical study',
            '시험대상자': 'subject',
            '피험자': 'subject',
            '시험책임자': 'principal investigator',
            '시험자': 'investigator',
            '의뢰자': 'sponsor',
            '임상시험 의뢰자': 'sponsor',
            
            # Investigational products
            '임상시험용 의약품': 'investigational product',
            '임상시험용 의료기기': 'investigational device',
            '시험약': 'investigational product',
            '대조약': 'comparator',
            '위약': 'placebo',
            '활성 대조약': 'active comparator',
            
            # Safety terminology
            '이상반응': 'adverse event',
            '이상사례': 'adverse event',
            '중대한 이상반응': 'serious adverse event',
            '약물이상반응': 'adverse drug reaction',
            '예상하지 못한 중대한 약물이상반응': 'suspected unexpected serious adverse reaction',
            
            # Study design
            '무작위배정': 'randomization',
            '무작위': 'randomized',
            '눈가림': 'blinding',
            '이중눈가림': 'double-blind',
            '단일눈가림': 'single-blind',
            '공개': 'open-label',
            
            # Criteria and endpoints
            '선정기준': 'inclusion criteria',
            '제외기준': 'exclusion criteria',
            '평가변수': 'endpoint',
            '일차 평가변수': 'primary endpoint',
            '이차 평가변수': 'secondary endpoint',
            '탐색적 평가변수': 'exploratory endpoint',
            '유효성': 'efficacy',
            '안전성': 'safety',
            
            # Study phases
            '스크리닝': 'screening',
            '기저선': 'baseline',
            '기저치': 'baseline',
            '기저시점': 'baseline',
            '방문': 'visit',
            '추적관찰': 'follow-up',
            '추적조사': 'follow-up',
            
            # Discontinuation
            '중도탈락': 'discontinuation',
            '조기종료': 'early termination',
            '철회': 'withdrawal',
            '동의철회': 'withdrawal of consent',
            
            # Documentation
            '임상시험 계획서': 'protocol',
            '증례기록서': 'case report form',
            '동의서': 'informed consent',
            '시험대상자 동의서': 'informed consent form',
            '설명문': 'subject information sheet',
            
            # Regulatory
            '의약품 임상시험 관리기준': 'Good Clinical Practice',
            '임상시험 관리기준': 'Good Clinical Practice',
            'GCP': 'GCP',
            'ICH': 'ICH',
            '임상시험심사위원회': 'Institutional Review Board',
            'IRB': 'IRB',
            
            # Analysis
            '통계분석': 'statistical analysis',
            '중간분석': 'interim analysis',
            '최종분석': 'final analysis',
            '분석군': 'analysis set',
            'ITT': 'intention-to-treat',
            'PP': 'per-protocol',
            
            # Time points
            '치료 종료': 'end of treatment',
            '임상시험 종료': 'end of study',
            'EOT': 'EOT',
            'EOS': 'EOS'
        }
        
        self.abbreviations = {
            'AE': 'adverse event',
            'SAE': 'serious adverse event',
            'ADR': 'adverse drug reaction',
            'SUSAR': 'suspected unexpected serious adverse reaction',
            'IP': 'investigational product',
            'CRF': 'case report form',
            'ICF': 'informed consent form',
            'PI': 'principal investigator',
            'GCP': 'Good Clinical Practice',
            'IRB': 'Institutional Review Board',
            'IEC': 'Independent Ethics Committee',
            'ITT': 'intention-to-treat',
            'PP': 'per-protocol',
            'FAS': 'full analysis set',
            'EOT': 'end of treatment',
            'EOS': 'end of study'
        }

    def extract_from_docx(self, ko_path: str, en_path: str) -> Dict:
        """Extract aligned terms from Korean and English protocol documents"""
        
        # Load documents
        doc_ko = Document(ko_path)
        doc_en = Document(en_path)
        
        # Extract all text
        ko_paragraphs = [p.text.strip() for p in doc_ko.paragraphs if p.text.strip()]
        en_paragraphs = [p.text.strip() for p in doc_en.paragraphs if p.text.strip()]
        
        # Extract from tables
        ko_table_text = []
        en_table_text = []
        
        for table in doc_ko.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        ko_table_text.append(cell.text.strip())
        
        for table in doc_en.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        en_table_text.append(cell.text.strip())
        
        # Combine all text
        all_ko_text = ko_paragraphs + ko_table_text
        all_en_text = en_paragraphs + en_table_text
        
        return {
            'korean_text': all_ko_text,
            'english_text': all_en_text,
            'paragraph_count': len(ko_paragraphs),
            'table_cell_count': len(ko_table_text)
        }
    
    def find_term_occurrences(self, text_list: List[str], terms: Dict[str, str]) -> Dict:
        """Find occurrences of terms in text"""
        
        occurrences = {}
        
        for ko_term, en_term in terms.items():
            ko_count = 0
            ko_contexts = []
            
            for text in text_list:
                if ko_term in text:
                    ko_count += text.count(ko_term)
                    if len(ko_contexts) < 3:  # Keep first 3 contexts
                        ko_contexts.append(text[:200])
            
            if ko_count > 0:
                occurrences[ko_term] = {
                    'english': en_term,
                    'count': ko_count,
                    'contexts': ko_contexts
                }
        
        return occurrences
    
    def extract_section_headers(self, text_list: List[str]) -> List[Tuple[str, str]]:
        """Extract numbered section headers"""
        
        headers = []
        section_pattern = re.compile(r'^(\d+(?:\.\d+)*)\s+(.+)$')
        
        for i, text in enumerate(text_list):
            match = section_pattern.match(text)
            if match:
                number = match.group(1)
                title = match.group(2)
                headers.append((number, title))
        
        return headers
    
    def analyze_style_patterns(self, en_text_list: List[str]) -> Dict:
        """Analyze English style patterns"""
        
        all_text = ' '.join(en_text_list)
        
        # Count modal verbs
        modal_counts = {
            'must': len(re.findall(r'\bmust\b', all_text, re.IGNORECASE)),
            'shall': len(re.findall(r'\bshall\b', all_text, re.IGNORECASE)),
            'should': len(re.findall(r'\bshould\b', all_text, re.IGNORECASE)),
            'will': len(re.findall(r'\bwill\b', all_text, re.IGNORECASE)),
            'may': len(re.findall(r'\bmay\b', all_text, re.IGNORECASE)),
            'can': len(re.findall(r'\bcan\b', all_text, re.IGNORECASE)),
            'could': len(re.findall(r'\bcould\b', all_text, re.IGNORECASE))
        }
        
        # Count passive voice indicators
        passive_patterns = [
            r'\bwill be\b',
            r'\bhas been\b',
            r'\bhave been\b',
            r'\bwas\b',
            r'\bwere\b',
            r'\bis being\b',
            r'\bare being\b'
        ]
        
        passive_count = sum(len(re.findall(pattern, all_text, re.IGNORECASE)) 
                          for pattern in passive_patterns)
        
        # Count formal expressions
        formal_expressions = {
            'prior to': len(re.findall(r'\bprior to\b', all_text, re.IGNORECASE)),
            'following': len(re.findall(r'\bfollowing\b', all_text, re.IGNORECASE)),
            'in accordance with': len(re.findall(r'\bin accordance with\b', all_text, re.IGNORECASE)),
            'pursuant to': len(re.findall(r'\bpursuant to\b', all_text, re.IGNORECASE)),
            'thereof': len(re.findall(r'\bthereof\b', all_text, re.IGNORECASE)),
            'herein': len(re.findall(r'\bherein\b', all_text, re.IGNORECASE))
        }
        
        return {
            'modal_verbs': modal_counts,
            'passive_voice_count': passive_count,
            'formal_expressions': formal_expressions,
            'total_words': len(all_text.split())
        }
    
    def create_comprehensive_glossary(self, 
                                    extracted_data: Dict,
                                    existing_glossary_path: Optional[str] = None) -> pd.DataFrame:
        """Create comprehensive glossary combining extracted and existing terms"""
        
        # Start with core terms
        glossary_data = []
        
        # Add core terms
        for ko_term, en_term in self.core_terms.items():
            glossary_data.append({
                'Korean': ko_term,
                'English': en_term,
                'Category': 'Core Protocol Term',
                'Source': 'Built-in'
            })
        
        # Add found terms from document
        if 'term_occurrences' in extracted_data:
            for ko_term, data in extracted_data['term_occurrences'].items():
                if ko_term not in self.core_terms:
                    glossary_data.append({
                        'Korean': ko_term,
                        'English': data['english'],
                        'Category': 'Document-specific',
                        'Source': 'Extracted'
                    })
        
        # Add existing glossary if provided
        if existing_glossary_path and Path(existing_glossary_path).exists():
            existing_df = pd.read_excel(existing_glossary_path)
            for _, row in existing_df.iterrows():
                glossary_data.append({
                    'Korean': row.get('KO', row.get('Korean', '')),
                    'English': row.get('EN', row.get('English', '')),
                    'Category': 'External Glossary',
                    'Source': 'Clinical Trial Reference'
                })
        
        # Create DataFrame and remove duplicates
        df = pd.DataFrame(glossary_data)
        df = df.drop_duplicates(subset=['Korean', 'English'])
        df = df.sort_values('Korean')
        
        return df
    
    def save_results(self, results: Dict, output_dir: str = '.'):
        """Save extraction results to files"""
        
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        
        # Save JSON data
        with open(output_path / 'protocol_terms.json', 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2)
        
        # Save glossary to Excel
        if 'glossary' in results:
            glossary_df = pd.DataFrame(results['glossary'])
            glossary_df.to_excel(output_path / 'protocol_glossary.xlsx', index=False)
        
        # Save style guide
        with open(output_path / 'style_guide.md', 'w', encoding='utf-8') as f:
            f.write("# Clinical Trial Protocol Style Guide\n\n")
            
            if 'style_analysis' in results:
                style = results['style_analysis']
                
                f.write("## Modal Verb Usage\n\n")
                for verb, count in style['modal_verbs'].items():
                    f.write(f"- **{verb}**: {count} occurrences\n")
                
                f.write(f"\n## Voice\n")
                f.write(f"- Passive voice indicators: {style['passive_voice_count']}\n")
                f.write(f"- Total words: {style['total_words']}\n")
                
                f.write("\n## Formal Expressions\n\n")
                for expr, count in style['formal_expressions'].items():
                    if count > 0:
                        f.write(f"- **{expr}**: {count} occurrences\n")
        
        print(f"Results saved to {output_path}")


def main():
    """Main execution function"""
    
    # Initialize extractor
    extractor = ProtocolTermExtractor()
    
    # Define paths
    base_path = Path("./Phase 2_AI testing kit/한영")
    ko_doc = base_path / "[원본 참고용] DW_DWP14012311_Protocol_v1.0_20230601.docx"
    en_doc = base_path / "[번역본 참고용] DW_DWP14012311_Protocol_v1.0_ENG_v1.0_20230706.docx"
    glossary_path = base_path / "2_용어집_SAMPLE_CLIENT KO-EN Clinical Trial Reference_20250421.xlsx"
    
    # Extract from documents
    print("Extracting terms from protocol documents...")
    doc_data = extractor.extract_from_docx(str(ko_doc), str(en_doc))
    
    # Find term occurrences
    print("Finding term occurrences...")
    term_occurrences = extractor.find_term_occurrences(
        doc_data['korean_text'], 
        extractor.core_terms
    )
    
    # Extract section headers
    print("Extracting section structure...")
    ko_headers = extractor.extract_section_headers(doc_data['korean_text'])
    en_headers = extractor.extract_section_headers(doc_data['english_text'])
    
    # Analyze style
    print("Analyzing English style patterns...")
    style_analysis = extractor.analyze_style_patterns(doc_data['english_text'])
    
    # Create comprehensive glossary
    print("Creating comprehensive glossary...")
    glossary_df = extractor.create_comprehensive_glossary(
        {'term_occurrences': term_occurrences},
        str(glossary_path) if glossary_path.exists() else None
    )
    
    # Compile results
    results = {
        'document_stats': {
            'paragraph_count': doc_data['paragraph_count'],
            'table_cell_count': doc_data['table_cell_count'],
            'total_korean_items': len(doc_data['korean_text']),
            'total_english_items': len(doc_data['english_text'])
        },
        'term_occurrences': term_occurrences,
        'style_analysis': style_analysis,
        'glossary': glossary_df.to_dict('records'),
        'section_headers': {
            'korean': [f"{num} {title}" for num, title in ko_headers[:20]],
            'english': [f"{num} {title}" for num, title in en_headers[:20]]
        }
    }
    
    # Save results
    output_dir = "./clinical_protocol_system/data"
    extractor.save_results(results, output_dir)
    
    # Print summary
    print("\n=== EXTRACTION SUMMARY ===")
    print(f"Documents analyzed: 2 (Korean source, English translation)")
    print(f"Paragraphs processed: {doc_data['paragraph_count']}")
    print(f"Table cells processed: {doc_data['table_cell_count']}")
    print(f"Terms identified: {len(term_occurrences)}")
    print(f"Glossary entries: {len(glossary_df)}")
    print(f"\nFiles created:")
    print(f"  - {output_dir}/protocol_terms.json")
    print(f"  - {output_dir}/protocol_glossary.xlsx")
    print(f"  - {output_dir}/style_guide.md")


if __name__ == "__main__":
    main()